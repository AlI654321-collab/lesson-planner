from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ایجاد یک فایل Word ساده برای تست
doc = Document()

# اضافه کردن عنوان
doc.add_heading('تست فایل Word', 0)

# اضافه کردن یک پاراگراف
paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = paragraph.add_run('این یک فایل Word تستی است.')

# ایجاد پوشه generated اگر وجود ندارد
if not os.path.exists("generated"):
    os.makedirs("generated")

# ذخیره فایل
doc.save("generated/test.docx")

print("فایل Word تستی با موفقیت تولید شد.")