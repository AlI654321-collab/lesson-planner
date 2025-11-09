from flask import Flask, request, jsonify, send_from_directory, render_template_string
import google.generativeai as genai
import os
import PyPDF2
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid

app = Flask(__name__, static_folder='.', template_folder='.')

# تنظیمات API
API_KEY = os.environ.get('GEMINI_API_KEY', 'AIzaSyCdRL9mQBAotXCLgyu_BNkaZVu_juL2yok')
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash')  # مدل جدید و سریع

# ذخیره محتوای فایل‌ها
syllabus_content = ""
book_content = ""

@app.route('/')
def index():
    try:
        with open('chatbot_new.html', 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error loading HTML: {e}")
        return """
        <html dir="rtl">
        <head><meta charset="UTF-8"><title>طرح درس ساز</title></head>
        <body style="font-family: Tahoma; padding: 20px; text-align: center;">
            <h1>🎓 طرح درس ساز آنلاین</h1>
            <p>خطا در بارگذاری صفحه: """ + str(e) + """</p>
            <p>لطفاً چند لحظه صبر کنید و صفحه را رفرش کنید.</p>
        </body>
        </html>
        """

@app.route('/chatbot_new.html')
def chatbot_new():
    return index()

@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({
        'status': 'ok',
        'message': 'سرویس فعال است',
        'api_key_set': bool(os.environ.get('GEMINI_API_KEY')),
        'files_uploaded': {
            'syllabus': bool(syllabus_content),
            'book': bool(book_content)
        }
    })

@app.route('/api/status')
def api_status():
    """وضعیت API"""
    global syllabus_content, book_content
    return jsonify({
        'status': 'ok',
        'syllabus_uploaded': len(syllabus_content) > 0,
        'book_uploaded': len(book_content) > 0,
        'syllabus_size': len(syllabus_content),
        'book_size': len(book_content),
        'api_key_configured': bool(os.environ.get('GEMINI_API_KEY'))
    })

@app.route('/test_ai')
def test_ai():
    """تست اتصال به Gemini AI"""
    try:
        print("🔌 تست اتصال به AI...")
        response = model.generate_content("سلام")
        print("✅ اتصال موفق!")
        return jsonify({
            'status': 'success',
            'message': 'اتصال به AI موفق بود'
        })
    except Exception as e:
        print(f"❌ خطا در اتصال: {e}")
        return jsonify({
            'status': 'error',
            'message': str(e)
        })

@app.route('/upload', methods=['POST'])
def upload_files():
    global syllabus_content, book_content
    
    try:
        print("📤 درخواست آپلود دریافت شد")
        
        if not request.files:
            return jsonify({
                'status': 'error',
                'message': 'هیچ فایلی ارسال نشده است'
            }), 400
        
        syllabus_name = None
        book_name = None
        book_info = {}
        
        if 'syllabus' in request.files:
            syllabus_file = request.files['syllabus']
            if syllabus_file and syllabus_file.filename != '':
                print(f"📄 در حال خواندن طرح درس نمونه: {syllabus_file.filename}")
                try:
                    syllabus_content = extract_text_from_pdf(syllabus_file)
                    syllabus_name = syllabus_file.filename
                    print(f"✓ طرح درس: {len(syllabus_content)} کاراکتر")
                except Exception as e:
                    print(f"✗ خطا در خواندن طرح درس: {e}")
                    return jsonify({
                        'status': 'error',
                        'message': f'خطا در خواندن فایل طرح درس: {str(e)}'
                    }), 400
        
        if 'book' in request.files:
            book_file = request.files['book']
            if book_file and book_file.filename != '':
                print(f"📚 در حال خواندن کل کتاب: {book_file.filename}")
                try:
                    book_content = extract_text_from_pdf(book_file)
                    book_name = book_file.filename
                    print(f"✓ کتاب: {len(book_content)} کاراکتر")
                    
                    # تحلیل خودکار کتاب
                    print("🔍 در حال تحلیل محتوای کتاب...")
                    try:
                        book_info = analyze_book_content(book_content)
                    except Exception as e:
                        print(f"⚠️ خطا در تحلیل کتاب: {e}")
                        book_info = {'course_name': 'نامشخص'}
                except Exception as e:
                    print(f"✗ خطا در خواندن کتاب: {e}")
                    return jsonify({
                        'status': 'error',
                        'message': f'خطا در خواندن فایل کتاب: {str(e)}'
                    }), 400
        
        return jsonify({
            'status': 'success',
            'syllabus_name': syllabus_name,
            'book_name': book_name,
            'book_info': book_info,
            'message': 'فایل‌ها با موفقیت پردازش شدند'
        })
    except Exception as e:
        print(f"✗ خطای کلی: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': f'خطای سرور: {str(e)}'
        }), 500

def analyze_book_content(content):
    """تحلیل خودکار محتوای کتاب"""
    try:
        print("  🤖 استفاده از AI برای تحلیل...")
        
        # استفاده از 15000 کاراکتر اول برای تحلیل بهتر
        sample_content = content[:15000]
        print(f"  📊 طول نمونه برای تحلیل: {len(sample_content)} کاراکتر")
        
        analysis_prompt = f"""
محتوای کتاب درسی (صفحات اول):
{sample_content}

لطفاً این اطلاعات را استخراج کنید و به صورت JSON پاسخ دهید:

{{
    "course_name": "نام کامل درس",
    "grade": "پایه تحصیلی",
    "field": "رشته تحصیلی",
    "chapters_count": تعداد فصل‌ها,
    "chapters": ["عنوان فصل 1", "عنوان فصل 2"],
    "suggested_request": "یک طرح درس کامل سالانه برای [نام درس] پایه [پایه] رشته [رشته] بساز"
}}

فقط JSON را برگردانید.
"""
        
        response = model.generate_content(analysis_prompt)
        result_text = response.text
        
        import json
        import re
        
        result_text = re.sub(r'```json\s*', '', result_text)
        result_text = re.sub(r'```\s*', '', result_text)
        
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if json_match:
            book_info = json.loads(json_match.group())
            print(f"  ✓ تحلیل کامل شد: {book_info.get('course_name', 'نامشخص')}")
            return book_info
        else:
            print("  ⚠️ JSON معتبر پیدا نشد")
            return {
                'course_name': 'نامشخص',
                'suggested_request': 'یک طرح درس کامل سالانه برای این کتاب بساز'
            }
    except Exception as e:
        print(f"  ⚠️ خطا در تحلیل: {e}")
        import traceback
        traceback.print_exc()
        return {
            'course_name': 'نامشخص',
            'suggested_request': 'یک طرح درس کامل سالانه برای این کتاب بساز'
        }

def extract_text_from_pdf(file):
    """خواندن کل محتوای PDF"""
    try:
        pdf_data = file.read()
        pdf_file = io.BytesIO(pdf_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        total = len(pdf_reader.pages)
        print(f"  📊 تعداد صفحات: {total}")
        
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            if (i + 1) % 20 == 0:
                print(f"  ⏳ خوانده شد: {i + 1}/{total}")
        
        print(f"  ✓ کل محتوای خوانده شده: {len(text)} کاراکتر")
        
        if len(text) < 100:
            print("  ⚠️ هشدار: محتوای خوانده شده خیلی کم است!")
            return "خطا: فایل PDF خالی است یا قابل خواندن نیست"
        
        return text
    except Exception as e:
        print(f"  ✗ خطا در خواندن PDF: {e}")
        import traceback
        traceback.print_exc()
        return f"خطا: {str(e)}"

def create_table_from_markdown(doc, table_lines):
    """تبدیل جدول markdown به Word با پشتیبانی فارسی"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    if not table_lines:
        return
    
    clean_lines = [line for line in table_lines 
                   if line.strip().replace('|', '').replace('-', '').strip()]
    
    if not clean_lines:
        return
    
    rows = []
    for line in clean_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < max_cols:
                cell = table.rows[i].cells[j]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                text = cell_data.replace('<br>', '\n')
                if '**' in text:
                    parts = text.split('**')
                    for k, part in enumerate(parts):
                        run = paragraph.add_run(part)
                        run.font.name = 'B Nazanin'
                        run.font.size = Pt(11)
                        run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                        if k % 2 == 1:
                            run.bold = True
                else:
                    run = paragraph.add_run(text)
                    run.font.name = 'B Nazanin'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
    
    doc.add_paragraph()

def split_book_into_chunks(content, chunk_size=80000):
    """تقسیم کتاب به بخش‌های کوچکتر"""
    chunks = []
    for i in range(0, len(content), chunk_size):
        chunks.append(content[i:i + chunk_size])
    return chunks

def extract_book_summary(book_content):
    """استخراج خلاصه کامل کتاب با پردازش چند مرحله‌ای"""
    print("📚 استخراج خلاصه کامل کتاب...")
    
    # اگر کتاب کوچک است، همه را برگردان
    if len(book_content) <= 100000:
        print(f"  ✓ کتاب کوچک است ({len(book_content)} کاراکتر) - استفاده از کل محتوا")
        return book_content
    
    # تقسیم به بخش‌ها
    chunks = split_book_into_chunks(book_content, 80000)
    print(f"  📊 کتاب به {len(chunks)} بخش تقسیم شد")
    
    summaries = []
    for i, chunk in enumerate(chunks):
        print(f"  ⏳ پردازش بخش {i+1}/{len(chunks)}...")
        
        prompt = f"""لطفاً این بخش از کتاب درسی را خلاصه کنید و فصل‌ها، موضوعات و مفاهیم کلیدی را استخراج کنید:

{chunk[:50000]}

خروجی: فهرست فصل‌ها و موضوعات به صورت خلاصه و ساختاریافته
"""
        
        try:
            response = model.generate_content(prompt)
            summaries.append(response.text)
            print(f"    ✓ بخش {i+1} پردازش شد")
        except Exception as e:
            print(f"    ⚠️ خطا در بخش {i+1}: {e}")
            summaries.append(chunk[:10000])  # استفاده از 10000 کاراکتر اول
    
    # ترکیب خلاصه‌ها
    combined_summary = "\n\n".join(summaries)
    print(f"  ✓ خلاصه کامل آماده شد: {len(combined_summary)} کاراکتر")
    
    return combined_summary

@app.route('/generate_word', methods=['POST'])
def generate_word():
    global syllabus_content, book_content
    
    try:
        # چک کردن فایل‌ها
        if not syllabus_content or not book_content:
            return jsonify({
                'status': 'error',
                'message': 'لطفاً ابتدا فایل طرح درس نمونه و کتاب را آپلود کنید'
            }), 400
        
        data = request.json
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'داده‌های ورودی خالی است'
            }), 400
            
        user_message = data.get('message', '')
        first_name = data.get('firstName', '')
        last_name = data.get('lastName', '')
        school_name = data.get('schoolName', '')
        class_day = data.get('classDay', 'شنبه')
        hours_per_week = data.get('hoursPerWeek', '8')
        holidays = data.get('holidays', '')
        
        print(f"\n{'='*60}")
        print(f"📝 درخواست: {user_message}")
        print(f"👤 نام: {first_name} {last_name}")
        print(f"🏫 مدرسه: {school_name}")
        # تعطیلات رسمی ایران
        holidays = """
تعطیلات رسمی سال تحصیلی:
- 22 بهمن: پیروزی انقلاب اسلامی (تعطیل)
- 29 اسفند: روز ملی شدن صنعت نفت (تعطیل)
- 1-4 فروردین: عید نوروز (تعطیل)
- 12-13 فروردین: روز طبیعت (تعطیل)
- 14 خرداد: رحلت امام خمینی (تعطیل)
- 15 خرداد: قیام 15 خرداد (تعطیل)
- تعطیلات مذهبی: شهادت امام علی، عید فطر، عید قربان، تاسوعا و عاشورا (تاریخ متغیر)
- تعطیلات نیمسال: 15 دی تا 24 دی (10 روز)
- پنجشنبه‌ها و جمعه‌ها: تعطیل هفتگی
"""
        
        print(f"📅 روز کلاس: {class_day}")
        print(f"⏰ ساعت: {hours_per_week}")
        print(f"📊 طول محتوای کتاب: {len(book_content)} کاراکتر")
        print(f"📊 طول طرح درس نمونه: {len(syllabus_content)} کاراکتر")
        print(f"{'='*60}")
        
        # استخراج خلاصه هوشمند از کل کتاب
        print("🔄 پردازش هوشمند کتاب...")
        book_summary = extract_book_summary(book_content)
        
        # محاسبه سال تحصیلی جاری
        import jdatetime
        today = jdatetime.datetime.now()
        # اگر ماه 1 تا 6 باشه (فروردین تا شهریور)، سال قبل شروع شده
        # اگر ماه 7 تا 12 باشه (مهر تا اسفند)، سال جاری شروع شده
        if today.month <= 6:
            current_year = today.year - 1
        else:
            current_year = today.year
        next_year = current_year + 1
        
        print(f"📅 سال تحصیلی: {current_year}-{next_year}")
        
        prompt = f"""شما متخصص طراحی طرح درس هستید.

طرح درس نمونه (ساختار دقیق - استفاده از کل محتوا):
{syllabus_content}

محتوای کامل کتاب درسی (پردازش شده):
{book_summary}

درخواست: {user_message}

اطلاعات معلم و مدرسه:
- نام معلم: {first_name} {last_name}
- نام مدرسه: {school_name}
- روز برگزاری کلاس: {class_day}
- تعداد ساعت در هفته: {hours_per_week} ساعت
- تعطیلات رسمی: {holidays if holidays else 'بدون تعطیلات خاص'}

دستورالعمل مهم:
1. در سرصفحه این اطلاعات را قید کنید:
   - نام معلم: {first_name} {last_name}
   - نام مدرسه: {school_name}
   - روز کلاس: {class_day}
   - ساعت هفتگی: {hours_per_week} ساعت
2. طرح درس را دقیقاً مثل نمونه بسازید
3. از تمام محتوای کتاب برای تعیین موضوعات و فصل‌ها استفاده کنید
4. همه ماه‌های سال تحصیلی را پوشش دهید: مهر، آبان، آذر، دی، بهمن، اسفند، فروردین، اردیبهشت، خرداد
5. برای هر {class_day} یک ردیف جدول بسازید (از اول مهر سال جاری تا آخر خرداد سال بعد)
6. فرمت جدول: | ماه | تاریخ کامل | ساعت | عنوان | اهداف | فعالیتها | توضیحات |
7. سال تحصیلی جاری: {current_year} (مهر تا اسفند) و {next_year} (فروردین تا خرداد)
8. در ستون "تاریخ کامل" فقط روز انتخابی ({class_day}) و تاریخ آن روز را بنویسید
9. فرمت دقیق: {class_day} سال/ماه/روز (مثلاً: {class_day} {current_year}/7/1)
10. مهم: روز انتخابی کاربر "{class_day}" است. فقط این روز را در جدول بیاورید.
11. روز انتخابی کاربر: {class_day}
12. اگر کاربر "دوشنبه" انتخاب کرده، تاریخ‌های دوشنبه را بنویسید (نه سه‌شنبه!)
13. در ستون "تاریخ کامل" روز هفته و تاریخ را بنویسید
14. فرمت: {class_day} سال/ماه/روز
15. تاریخ‌های {class_day} در سال {current_year}:
    - اگر دوشنبه: مهر 7، 14، 21، 28 (یعنی {current_year}/7/7، {current_year}/7/14، {current_year}/7/21، {current_year}/7/28)
    - اگر سه‌شنبه: مهر 1، 8، 15، 22، 29 (یعنی {current_year}/7/1، {current_year}/7/8، {current_year}/7/15، {current_year}/7/22، {current_year}/7/29)
    - اگر چهارشنبه: مهر 2، 9، 16، 23، 30 (یعنی {current_year}/7/2، {current_year}/7/9، {current_year}/7/16، {current_year}/7/23، {current_year}/7/30)
16. مثال برای دوشنبه:
    | مهر | دوشنبه {current_year}/7/7 | 8 | فصل اول | ...
    | مهر | دوشنبه {current_year}/7/14 | 8 | ادامه | ...
    | مهر | دوشنبه {current_year}/7/21 | 8 | فصل دوم | ...
17. مثال برای سه‌شنبه:
    | مهر | سه‌شنبه {current_year}/7/1 | 8 | فصل اول | ...
    | مهر | سه‌شنبه {current_year}/7/8 | 8 | ادامه | ...
18. در ستون ساعت، از {hours_per_week} ساعت استفاده کنید
19. مهم: حتماً روز انتخابی کاربر ({class_day}) را در نظر بگیرید!
9. تعطیلات رسمی را در نظر بگیرید و در آن روزها هیچ برنامه‌ای ننویسید
10. برای روزهای تعطیل، فقط یک ردیف خالی با "تعطیل رسمی - [نام تعطیلات]" در ستون توضیحات بگذارید
11. تعطیلات نیمسال (15-24 دی) را کاملاً خالی بگذارید
12. جدول باید کامل و جامع باشد و تمام فصل‌های کتاب را پوشش دهد

خروجی باید شامل:
- بدون هیچ مقدمه یا توضیح اضافی، مستقیماً با سرصفحه شروع کنید
- سرصفحه (نام درس، پایه، رشته، نام معلم، نام مدرسه، روز کلاس: {class_day}، ساعت هفتگی، سال تحصیلی {current_year}-{next_year})
- جدول کامل برنامه سالانه با روز هفته و تاریخ
- فرمت ستون تاریخ: {class_day} سال/ماه/روز
- اگر کاربر دوشنبه انتخاب کرده، مثال ردیف‌ها:
    | مهر | دوشنبه {current_year}/7/7 | 8 | فصل اول | ...
    | مهر | دوشنبه {current_year}/7/14 | 8 | ادامه | ...
    | مهر | دوشنبه {current_year}/7/21 | 8 | فصل دوم | ...
- اگر کاربر سه‌شنبه انتخاب کرده، مثال ردیف‌ها:
    | مهر | سه‌شنبه {current_year}/7/1 | 8 | فصل اول | ...
    | مهر | سه‌شنبه {current_year}/7/8 | 8 | ادامه | ...
- تعطیلات نیمسال (15-24 دی) را کاملاً خالی بگذارید
- ابزارها و وسایل آموزشی
- روش‌های ارزشیابی
- منابع و مراجع

مهم: 
1. هیچ جمله توضیحی مانند "بسیار خب، این یک طرح درس..." ننویسید
2. حتماً روز انتخابی کاربر ({class_day}) را در جدول بنویسید
3. اگر دوشنبه انتخاب شده، تاریخ‌های دوشنبه بنویسید (7، 14، 21، 28)
4. اگر سه‌شنبه انتخاب شده، تاریخ‌های سه‌شنبه بنویسید (1، 8، 15، 22، 29)
"""

        print("🤖 در حال ارسال به Gemini AI...")
        response = model.generate_content(prompt)
        syllabus_text = response.text
        print(f"✓ پاسخ دریافت شد: {len(syllabus_text)} کاراکتر")
        
        print("📄 در حال ایجاد فایل Word...")
        doc = Document()
        
        lines = syllabus_text.split('\n')
        in_table = False
        table_lines = []
        
        for line in lines:
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
            else:
                if in_table and table_lines:
                    create_table_from_markdown(doc, table_lines)
                    in_table = False
                    table_lines = []
                
                if line.strip():
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    if '**' in line:
                        parts = line.split('**')
                        for i, part in enumerate(parts):
                            run = paragraph.add_run(part)
                            if i % 2 == 1:
                                run.bold = True
                    else:
                        paragraph.add_run(line)
        
        if in_table and table_lines:
            create_table_from_markdown(doc, table_lines)
        
        filename = f"طرح_درس_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join("generated", filename)
        
        if not os.path.exists("generated"):
            os.makedirs("generated")
        
        doc.save(filepath)
        print(f"✓ فایل ذخیره شد: {filename}\n")
        
        return jsonify({
            'status': 'success',
            'content': syllabus_text,
            'filename': filename,
            'message': 'طرح درس با موفقیت تولید شد'
        })
    except Exception as e:
        print(f"✗ خطا: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/export', methods=['POST'])
def export_file():
    """تولید فایل با فرمت دلخواه"""
    try:
        data = request.json
        content = data.get('content', '')
        format_type = data.get('format', 'word')
        
        print(f"📤 در حال تولید فایل {format_type.upper()}...")
        
        filename_base = f"طرح_درس_{uuid.uuid4().hex[:8]}"
        
        if format_type == 'pdf':
            filename = f"{filename_base}.pdf"
            filepath = create_pdf_file(content, filename)
        elif format_type == 'word':
            filename = f"{filename_base}.docx"
            filepath = create_word_file(content, filename)
        elif format_type == 'excel':
            filename = f"{filename_base}.xlsx"
            filepath = create_excel_file(content, filename)
        elif format_type == 'html':
            filename = f"{filename_base}.html"
            filepath = create_html_file(content, filename)
        else:
            return jsonify({'status': 'error', 'message': 'فرمت نامعتبر'})
        
        print(f"✓ فایل ذخیره شد: {filename}")
        
        return jsonify({
            'status': 'success',
            'filename': filename,
            'message': f'فایل {format_type.upper()} با موفقیت تولید شد'
        })
    except Exception as e:
        print(f"✗ خطا: {e}")
        return jsonify({'status': 'error', 'message': str(e)})

def create_word_file(content, filename):
    """ایجاد فایل Word با پشتیبانی کامل از فارسی"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # تنظیم فونت پیش‌فرض برای فارسی
    style = doc.styles['Normal']
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(12)
    
    # تنظیم فونت برای زبان‌های پیچیده (فارسی/عربی)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    rFonts.set(qn('w:cs'), 'B Nazanin')
    
    lines = content.split('\n')
    in_table = False
    table_lines = []
    
    for line in lines:
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        else:
            if in_table and table_lines:
                create_table_from_markdown(doc, table_lines)
                in_table = False
                table_lines = []
            
            if line.strip():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = paragraph.add_run(part)
                        run.font.name = 'B Nazanin'
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                        if i % 2 == 1:
                            run.bold = True
                else:
                    run = paragraph.add_run(line)
                    run.font.name = 'B Nazanin'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
    
    if in_table and table_lines:
        create_table_from_markdown(doc, table_lines)
    
    filepath = os.path.join("generated", filename)
    if not os.path.exists("generated"):
        os.makedirs("generated")
    
    doc.save(filepath)
    return filepath

def create_excel_file(content, filename):
    """ایجاد فایل Excel"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "طرح درس"
    
    lines = content.split('\n')
    row = 1
    
    for line in lines:
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            for col, cell_data in enumerate(cells, 1):
                cell = ws.cell(row=row, column=col, value=cell_data)
                cell.alignment = Alignment(horizontal='right', vertical='top', wrap_text=True)
                
                if row == 1:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            row += 1
        elif line.strip() and not line.strip().replace('-', ''):
            continue
        elif line.strip():
            ws.cell(row=row, column=1, value=line.strip())
            row += 1
    
    filepath = os.path.join("generated", filename)
    wb.save(filepath)
    return filepath

def create_html_file(content, filename):
    """ایجاد فایل HTML"""
    import markdown
    html_content = markdown.markdown(content, extensions=['tables'])
    
    html_template = f"""<!DOCTYPE html>
<html dir="rtl" lang="fa">
<head>
    <meta charset="UTF-8">
    <title>طرح درس</title>
    <style>
        body {{ font-family: Tahoma; padding: 20px; direction: rtl; text-align: right; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: right; }}
        th {{ background: #4472C4; color: white; }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""
    
    filepath = os.path.join("generated", filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html_template)
    
    return filepath

def create_pdf_file(content, filename):
    """ایجاد فایل PDF از HTML با پشتیبانی فارسی"""
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        # ایجاد HTML با فونت فارسی
        html_content = f"""<!DOCTYPE html>
<html dir="rtl" lang="fa">
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: 'Tahoma', 'Arial', sans-serif;
            direction: rtl;
            text-align: right;
            line-height: 1.8;
            font-size: 12pt;
        }}
        h1, h2, h3 {{
            color: #2c3e50;
            margin: 15px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            font-size: 10pt;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: right;
        }}
        th {{
            background-color: #4472C4;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
        .header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #4472C4;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>طرح درس سالانه</h1>
    </div>
    <pre style="white-space: pre-wrap; font-family: Tahoma;">{content}</pre>
</body>
</html>"""
        
        filepath = os.path.join("generated", filename)
        
        # تنظیمات فونت
        font_config = FontConfiguration()
        
        # تبدیل HTML به PDF
        HTML(string=html_content).write_pdf(
            filepath,
            font_config=font_config
        )
        
        return filepath
        
    except ImportError:
        # اگر weasyprint نصب نیست، از روش جایگزین استفاده کن
        print("⚠️ weasyprint نصب نیست، استفاده از روش جایگزین...")
        return create_pdf_from_word(content, filename)

def create_pdf_from_word(content, filename):
    """ایجاد PDF از طریق Word (روش جایگزین)"""
    try:
        # ابتدا Word بساز
        word_filename = filename.replace('.pdf', '.docx')
        word_filepath = create_word_file(content, word_filename)
        
        # سعی کن Word را به PDF تبدیل کنی
        try:
            from docx2pdf import convert
            pdf_filepath = word_filepath.replace('.docx', '.pdf')
            convert(word_filepath, pdf_filepath)
            return pdf_filepath
        except:
            # اگر نشد، فقط فایل Word را برگردان
            print("⚠️ تبدیل به PDF ممکن نیست، فایل Word ایجاد شد")
            return word_filepath
            
    except Exception as e:
        print(f"✗ خطا در ایجاد PDF: {e}")
        # در نهایت HTML بساز
        return create_html_file(content, filename.replace('.pdf', '.html'))

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_from_directory('generated', filename, as_attachment=True)
    except Exception as e:
        return str(e), 404

if __name__ == '__main__':
    os.makedirs('generated', exist_ok=True)
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
